import os
import io
import re
import time
import json
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from fastapi import HTTPException, UploadFile
import PyPDF2
import docx

# Set up logging
logger = logging.getLogger(__name__)

@dataclass
class SensitiveDataPattern:
    """Represents a pattern for detecting sensitive data"""
    pattern_type: str  # "ssn", "phone", "email", "credit_card", "custom"
    regex_pattern: str
    description: str
    risk_level: str  # "high", "medium", "low"

@dataclass
class SensitiveDataMatch:
    """Enhanced sensitive data detection with location information"""
    pattern_type: str
    description: str
    risk_level: str
    original_text: str
    start_position: int
    end_position: int
    context_before: str
    context_after: str
    page_number: Optional[int] = None
    line_number: Optional[int] = None

class ProcessedDocument:
    """Represents a processed document with its metadata"""
    def __init__(self, filename: str, content: str, sensitive_data_found: List[SensitiveDataMatch], 
                 processing_summary: Dict[str, Any], anonymized_content: str, content_tokens: List[Dict[str, Any]] = None, page_breakdown: Optional[List[Dict[str, Any]]] = None):
        self.filename = filename
        self.content = content
        self.sensitive_data_found = sensitive_data_found
        self.processing_summary = processing_summary
        self.anonymized_content = anonymized_content
        self.content_tokens = content_tokens or []
        self.page_breakdown = page_breakdown

class DocumentProcessor:
    """Main class for processing documents and detecting sensitive data"""
    
    def __init__(self):
        # Predefined sensitive data patterns
        self.patterns = [
            SensitiveDataPattern(
                pattern_type="ssn",
                regex_pattern=r"\b\d{3}-\d{2}-\d{4}\b",
                description="Social Security Number",
                risk_level="high"
            ),
            SensitiveDataPattern(
                pattern_type="phone",
                regex_pattern=r"\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b",
                description="Phone Number",
                risk_level="medium"
            ),
            SensitiveDataPattern(
                pattern_type="email",
                regex_pattern=r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
                description="Email Address",
                risk_level="medium"
            ),
            SensitiveDataPattern(
                pattern_type="credit_card",
                regex_pattern=r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b",
                description="Credit Card Number",
                risk_level="high"
            ),
            SensitiveDataPattern(
                pattern_type="address",
                regex_pattern=r"\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Drive|Dr|Lane|Ln|Way|Court|Ct)\b",
                description="Street Address",
                risk_level="medium"
            ),
            SensitiveDataPattern(
                pattern_type="name",
                regex_pattern=r"\b(?:Mr\.|Mrs\.|Ms\.|Dr\.)\s+[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b",
                description="Full Name",
                risk_level="medium"
            ),
            SensitiveDataPattern(
                pattern_type="patient_name",
                regex_pattern=r"\bPatient Name:\s*([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b",
                description="Patient Name",
                risk_level="high"
            ),
            SensitiveDataPattern(
                pattern_type="physician_name",
                regex_pattern=r"\bPhysician:\s*(?:Dr\.)?\s*([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b",
                description="Physician Name",
                risk_level="medium"
            ),
            SensitiveDataPattern(
                pattern_type="date_of_birth",
                regex_pattern=r"\b(?:0[1-9]|1[0-2])/(?:0[1-9]|[12]\d|3[01])/\d{4}\b",
                description="Date of Birth",
                risk_level="high"
            )
        ]
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            try:
                return file_content.decode('latin-1').strip()
            except Exception as e:
                logger.error(f"Error decoding text file: {e}")
                return ""
    
    def detect_sensitive_data(self, text: str) -> List[SensitiveDataPattern]:
        """Detect sensitive data patterns in text"""
        found_patterns = []
        
        for pattern in self.patterns:
            matches = re.finditer(pattern.regex_pattern, text, re.IGNORECASE)
            if matches:
                # Create a copy with actual matches found
                pattern_copy = SensitiveDataPattern(
                    pattern_type=pattern.pattern_type,
                    regex_pattern=pattern.regex_pattern,
                    description=pattern.description,
                    risk_level=pattern.risk_level
                )
                found_patterns.append(pattern_copy)
        
        return found_patterns
    
    def detect_sensitive_data_with_locations(self, text: str, file_type: str = "txt") -> List[SensitiveDataMatch]:
        """Detect sensitive data patterns with detailed location information"""
        found_matches = []
        
        for pattern in self.patterns:
            matches = re.finditer(pattern.regex_pattern, text, re.IGNORECASE)
            for match in matches:
                start_pos = match.start()
                end_pos = match.end()
                original_text = match.group()
                
                # Handle capture groups for patterns like patient_name
                if pattern.pattern_type in ["patient_name", "physician_name"]:
                    # Extract just the name part from the capture group
                    name_match = re.search(r"([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)", original_text)
                    if name_match:
                        original_text = name_match.group(1)
                        # Adjust positions to just the name part
                        start_pos = start_pos + name_match.start(1)
                        end_pos = start_pos + len(original_text)
                
                # Get context (50 characters before and after)
                context_before = text[max(0, start_pos - 50):start_pos]
                context_after = text[end_pos:min(len(text), end_pos + 50)]
                
                # Calculate line number for text-based documents
                line_number = text[:start_pos].count('\n') + 1 if file_type != "pdf" else None
                
                # Calculate page number for PDFs (approximate based on character count)
                page_number = None
                if file_type == "pdf":
                    # Assume average 2000 characters per page
                    page_number = (start_pos // 2000) + 1
                
                sensitive_match = SensitiveDataMatch(
                    pattern_type=pattern.pattern_type,
                    description=pattern.description,
                    risk_level=pattern.risk_level,
                    original_text=original_text,
                    start_position=start_pos,
                    end_position=end_pos,
                    page_number=page_number,
                    line_number=line_number,
                    context_before=context_before,
                    context_after=context_after
                )
                found_matches.append(sensitive_match)
        
        return found_matches
    
    def anonymize_text(self, text: str, anonymization_type: str = "redact") -> str:
        """Anonymize sensitive data in text"""
        anonymized_text = text
        
        for pattern in self.patterns:
            if anonymization_type == "redact":
                anonymized_text = re.sub(
                    pattern.regex_pattern, 
                    f"[{pattern.pattern_type.upper()}_REDACTED]", 
                    anonymized_text, 
                    flags=re.IGNORECASE
                )
            elif anonymization_type == "mask":
                anonymized_text = re.sub(
                    pattern.regex_pattern, 
                    "*" * 10, 
                    anonymized_text, 
                    flags=re.IGNORECASE
                )
            elif anonymization_type == "anonymize":
                anonymized_text = re.sub(
                    pattern.regex_pattern, 
                    self._generate_fake_data(pattern.pattern_type), 
                    anonymized_text, 
                    flags=re.IGNORECASE
                )
        
        return anonymized_text
    
    def _generate_fake_data(self, pattern_type: str) -> str:
        """Generate fake data for anonymization"""
        fake_data_map = {
            "ssn": "XXX-XX-XXXX",
            "phone": "(555) 123-4567",
            "email": "user@example.com",
            "credit_card": "****-****-****-1234",
            "address": "123 Main Street",
            "name": "John Doe",
            "date_of_birth": "01/01/1990"
        }
        return fake_data_map.get(pattern_type, "[ANONYMIZED]")
    
    def process_single_document(self, file: UploadFile, anonymization_type: str = "redact") -> ProcessedDocument:
        """Process a single document and return processed data"""
        if not file.filename:
            raise ValueError("File must have a filename")
        
        # Read file content
        content = file.file.read()
        
        # Extract text based on file type
        text = ""
        filename_lower = file.filename.lower()
        
        if filename_lower.endswith('.pdf'):
            text = self.extract_text_from_pdf(content)
        elif filename_lower.endswith('.docx'):
            text = self.extract_text_from_docx(content)
        elif filename_lower.endswith('.txt'):
            text = self.extract_text_from_txt(content)
        else:
            raise ValueError(f"Unsupported file type: {file.filename}")
        
        if not text:
            raise ValueError(f"Could not extract text from {file.filename}")
        
        # Get file type for enhanced detection
        file_type = filename_lower.split('.')[-1]
        
        # Detect sensitive data with location information
        sensitive_data = self.detect_sensitive_data_with_locations(text, file_type)
        
        # Generate content tokens for frontend rendering
        content_tokens = self.generate_content_tokens(text, sensitive_data)
        
        # Generate page breakdown (for PDFs)
        page_breakdown = self.generate_page_breakdown(text, sensitive_data, file_type)
        
        # Anonymize content
        anonymized_content = self.anonymize_text(text, anonymization_type)
        
        # Create processing summary
        processing_summary = {
            "file_size": len(content),
            "text_length": len(text),
            "sensitive_patterns_found": len(sensitive_data),
            "anonymization_applied": anonymization_type,
            "processing_timestamp": datetime.now().isoformat(),
            "file_type": file_type
        }
        
        return ProcessedDocument(
            filename=file.filename,
            content=text,
            sensitive_data_found=sensitive_data,
            processing_summary=processing_summary,
            anonymized_content=anonymized_content,
            content_tokens=content_tokens,
            page_breakdown=page_breakdown
        )
    
    def process_multiple_documents(self, files: List[UploadFile], 
                                 anonymization_type: str = "redact") -> List[ProcessedDocument]:
        """Process multiple documents"""
        processed_docs = []
        
        for file in files:
            try:
                processed_doc = self.process_single_document(file, anonymization_type)
                processed_docs.append(processed_doc)
            except Exception as e:
                logger.error(f"Error processing {file.filename}: {e}")
                # Continue with other files instead of failing completely
                continue
        
        return processed_docs

    def generate_page_breakdown(self, text: str, sensitive_matches: List[SensitiveDataMatch], file_type: str) -> Optional[List[Dict[str, Any]]]:
        """Generate page-by-page breakdown of sensitive data"""
        if file_type != "pdf":
            return None
        
        # Group sensitive data by approximate page
        page_data = {}
        
        for match in sensitive_matches:
            page_num = match.page_number or 1
            if page_num not in page_data:
                page_data[page_num] = {
                    "page_number": page_num,
                    "sensitive_items": [],
                    "risk_summary": {"high": 0, "medium": 0, "low": 0}
                }
            
            page_data[page_num]["sensitive_items"].append({
                "type": match.pattern_type,
                "description": match.description,
                "risk_level": match.risk_level,
                "context": f"{match.context_before}...{match.context_after}"
            })
            
            page_data[page_num]["risk_summary"][match.risk_level] += 1
        
        return list(page_data.values())
    
    def generate_content_tokens(self, text: str, sensitive_matches: List[SensitiveDataMatch]) -> List[Dict[str, Any]]:
        """Generate tokenized content for frontend rendering"""
        if not sensitive_matches:
            return [{"type": "text", "content": text}]
        
        # Sort matches by position
        sorted_matches = sorted(sensitive_matches, key=lambda x: x.start_position)
        
        tokens = []
        last_pos = 0
        
        for match in sorted_matches:
            # Add text before sensitive data
            if match.start_position > last_pos:
                tokens.append({
                    "type": "text",
                    "content": text[last_pos:match.start_position]
                })
            
            # Add sensitive data token
            tokens.append({
                "type": "sensitive_data",
                "content": match.original_text,
                "pattern_type": match.pattern_type,
                "risk_level": match.risk_level,
                "description": match.description,
                "context_before": match.context_before,
                "context_after": match.context_after,
                "start_position": match.start_position,
                "end_position": match.end_position,
                "page_number": match.page_number,
                "line_number": match.line_number
            })
            
            last_pos = match.end_position
        
        # Add remaining text
        if last_pos < len(text):
            tokens.append({
                "type": "text",
                "content": text[last_pos:]
            })
        
        return tokens

# Initialize global document processor instance
document_processor = DocumentProcessor()
